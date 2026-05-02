"""
Generate a printable review document for partner sign-off.

Produces a Word .docx containing a side-by-side table (Original | Proposed)
for every flagged section. Partners can print or mark up this document
independently of the source file.
"""

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.models import ChangeType, ReviewState, Section


def generate_review_document(
    sections: list[Section],
    original_filename: str,
) -> tuple[bytes, str]:
    """
    Return (docx_bytes, output_filename).
    Only includes sections flagged for change (UPDATE or GAP).
    """
    doc = Document()

    # ── Page margins — slightly narrower to give tables more room ─────────────
    for section in doc.sections:
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Cover block ───────────────────────────────────────────────────────────
    title = doc.add_heading("Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run("Source document: ").bold = True
    meta.add_run(original_filename)

    meta2 = doc.add_paragraph()
    meta2.add_run("Date: ").bold = True
    meta2.add_run(date.today().strftime("%B %d, %Y"))

    flagged = [s for s in sections if s.change_type != ChangeType.NO_CHANGE]
    updates  = sum(1 for s in flagged if s.change_type == ChangeType.UPDATE)
    gaps     = sum(1 for s in flagged if s.change_type == ChangeType.GAP)

    meta3 = doc.add_paragraph()
    meta3.add_run("Sections flagged for review: ").bold = True
    meta3.add_run(f"{len(flagged)} ({updates} update(s), {gaps} gap(s))")

    doc.add_paragraph(
        "Instructions: Review each section below. Tick Accept or Reject for each proposed change "
        "and return this document to the document owner."
    ).italic = True

    if not flagged:
        doc.add_paragraph("No sections required changes.")
        return _save(doc, original_filename)

    doc.add_page_break()

    # ── One section per flagged item ──────────────────────────────────────────
    for i, section in enumerate(flagged, 1):
        change_label = "Update Needed" if section.change_type == ChangeType.UPDATE else "Gap Identified"

        heading = doc.add_heading(f"{i}. {section.heading}", level=1)

        info = doc.add_paragraph()
        run = info.add_run(f"Change type: {change_label}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x6C, 0x00) if section.change_type == ChangeType.UPDATE else RGBColor(0x17, 0x5C, 0xB0)

        if section.classify_reason:
            reason_p = doc.add_paragraph()
            reason_p.add_run("AI assessment: ").bold = True
            reason_p.add_run(section.classify_reason).italic = True

        # Side-by-side table
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        # Column widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.1)

        # Header row
        hdr = table.rows[0].cells
        _set_cell_header(hdr[0], "Original")
        _set_cell_header(hdr[1], "Proposed")

        # Content row
        content = table.rows[1].cells
        content[0].text = section.original_text or ""
        _set_cell_font(content[0])

        proposed_text = (
            section.final_text
            if section.review_state == ReviewState.EDITED and section.final_text
            else section.proposed_text or ""
        )
        content[1].text = proposed_text
        _set_cell_font(content[1])

        # Review checkbox row
        doc.add_paragraph("")
        checkbox_p = doc.add_paragraph()
        checkbox_p.add_run("Decision: ").bold = True
        checkbox_p.add_run("  ☐ Accept      ☐ Reject      ☐ Accept with edits")

        notes_p = doc.add_paragraph()
        notes_p.add_run("Notes: ").bold = True
        notes_p.add_run("_" * 80)

        doc.add_paragraph("")  # spacing between sections

    return _save(doc, original_filename)


def _set_cell_header(cell, text: str) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    # Light grey background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tc.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "D9D9D9",
    })
    tcPr.append(shd)


def _set_cell_font(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)


def _save(doc: Document, original_filename: str) -> tuple[bytes, str]:
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    stem  = Path(original_filename).stem
    today = date.today().strftime("%Y-%m-%d")
    return output.getvalue(), f"{stem}_review_{today}.docx"
