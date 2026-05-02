"""
Demo mode: generates in-memory sample documents so the app can be showcased
without any file uploads. Call get_demo_files() to receive the bytes for the
outdated document and the reference document.
"""

import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_outdated_document() -> bytes:
    """
    Returns bytes for a .docx representing an outdated product brief.
    Product name: 'NovaPay Pro'  (about to be renamed 'NovaPay Platform')
    Features, pricing, and SLA are intentionally stale.
    """
    doc = Document()

    doc.add_heading("NovaPay Pro — Product Brief", 0)

    _add_heading(doc, "Product Overview")
    _add_para(doc, (
        "NovaPay Pro is a cloud-based payment processing solution designed for "
        "small and medium-sized businesses. Launched in 2021, NovaPay Pro streamlines "
        "invoicing, recurring billing, and one-time payment collection through a "
        "simple web dashboard. The product targets finance teams that need a "
        "no-code tool to manage their payment workflows without engineering support."
    ))

    _add_heading(doc, "Core Features")
    _add_para(doc, (
        "NovaPay Pro includes the following features:\n"
        "• Invoice Builder: Create and send branded invoices in minutes.\n"
        "• Recurring Billing: Set up subscription billing with monthly or annual cycles.\n"
        "• Payment Links: Share a URL that customers use to pay directly.\n"
        "• Basic Reporting: View payment history and export CSV statements.\n"
        "• Email Notifications: Automatic reminders sent to customers for upcoming and overdue invoices."
    ))

    _add_heading(doc, "Supported Integrations")
    _add_para(doc, (
        "NovaPay Pro integrates with the following platforms out of the box:\n"
        "• QuickBooks Online\n"
        "• Xero\n"
        "• Stripe (as the underlying payment gateway)\n"
        "• Zapier (for custom workflow automation)\n"
        "All integrations are configured through the Settings → Integrations panel."
    ))

    _add_heading(doc, "Pricing")
    _add_para(doc, (
        "NovaPay Pro is available on three plans:\n\n"
        "Starter — $29 / month\n"
        "Up to 50 invoices per month, 1 user seat, email support.\n\n"
        "Growth — $79 / month\n"
        "Unlimited invoices, 5 user seats, priority email support, recurring billing.\n\n"
        "Business — $149 / month\n"
        "Everything in Growth, 20 user seats, custom branding, QuickBooks and Xero integration.\n\n"
        "All plans include a 14-day free trial. No credit card required to start."
    ))

    _add_heading(doc, "Security & Compliance")
    _add_para(doc, (
        "NovaPay Pro is PCI-DSS Level 2 compliant. Payment card data is never stored "
        "on NovaPay servers — all card processing is handled by Stripe. Customer data "
        "is encrypted at rest (AES-256) and in transit (TLS 1.2). SOC 2 Type I audit "
        "was completed in Q3 2022."
    ))

    _add_heading(doc, "Support & SLA")
    _add_para(doc, (
        "All NovaPay Pro customers receive email-based support with a 48-hour response "
        "time guarantee. Business plan customers are upgraded to 24-hour response. "
        "Support is available Monday through Friday, 9 AM – 5 PM EST. "
        "There is no phone or live-chat support available at this time. "
        "A self-service knowledge base is available at help.novapay.io."
    ))

    _add_heading(doc, "About NovaPay")
    _add_para(doc, (
        "NovaPay was founded in 2019 and is headquartered in Austin, Texas. "
        "The company processes over $200 million in payments annually and serves "
        "more than 3,000 business customers across North America. NovaPay is "
        "privately held and backed by Series A funding."
    ))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_reference_document() -> bytes:
    """
    Returns bytes for a .docx representing updated internal specs.
    Covers the NovaPay rebrand, new features, new pricing, and updated SLA.
    """
    doc = Document()

    doc.add_heading("NovaPay Internal — Q1 2025 Product Update", 0)

    _add_heading(doc, "Rebrand Notice")
    _add_para(doc, (
        "Effective January 2025, 'NovaPay Pro' is officially renamed 'NovaPay Platform'. "
        "All external collateral, website copy, and sales materials must reflect this change. "
        "The old name should not appear in any customer-facing documents after March 1, 2025."
    ))

    _add_heading(doc, "New Features — Q4 2024 / Q1 2025 Releases")
    _add_para(doc, (
        "The following features have been added since the last product brief was published:\n\n"
        "• AI Smart Reminders: The system now auto-schedules follow-up reminders based on "
        "each customer's payment history, reducing late payments by an average of 34%.\n\n"
        "• Multi-Currency Support: NovaPay Platform now supports 40+ currencies with real-time "
        "FX conversion. Available on Growth and Enterprise plans.\n\n"
        "• Mobile App (iOS & Android): A native mobile app launched in November 2024 allows "
        "users to send invoices, track payments, and receive push notifications on the go.\n\n"
        "• Bulk Invoice Import: Users can now upload a CSV to generate up to 500 invoices "
        "in a single batch operation.\n\n"
        "• Salesforce Integration: A native Salesforce connector is now available on "
        "the Enterprise plan, syncing deals and contacts automatically."
    ))

    _add_heading(doc, "Updated Integrations")
    _add_para(doc, (
        "The integrations lineup has expanded. Current supported integrations:\n"
        "• QuickBooks Online\n"
        "• Xero\n"
        "• FreshBooks (new — added Q4 2024)\n"
        "• Stripe (payment gateway)\n"
        "• PayPal (new — added Q1 2025)\n"
        "• Salesforce (Enterprise plan only)\n"
        "• HubSpot CRM (new — added Q1 2025)\n"
        "• Zapier\n"
        "• Make (formerly Integromat) (new — added Q1 2025)"
    ))

    _add_heading(doc, "2025 Pricing Update")
    _add_para(doc, (
        "Pricing has been updated effective February 1, 2025. "
        "Existing customers on legacy plans are grandfathered until their next annual renewal.\n\n"
        "Starter — $39 / month (was $29)\n"
        "Up to 100 invoices per month, 1 user seat, email support, mobile app access.\n\n"
        "Growth — $99 / month (was $79)\n"
        "Unlimited invoices, 10 user seats (was 5), priority email support, recurring billing, "
        "multi-currency, AI Smart Reminders.\n\n"
        "Business — $199 / month (was $149)\n"
        "Everything in Growth, 50 user seats (was 20), custom branding, all integrations, "
        "bulk invoice import, dedicated onboarding call.\n\n"
        "Enterprise — Custom pricing (new tier)\n"
        "Unlimited seats, Salesforce integration, SLA guarantees, SSO, custom contracts.\n\n"
        "Free trial extended to 21 days (was 14) across all plans."
    ))

    _add_heading(doc, "Security & Compliance Updates")
    _add_para(doc, (
        "NovaPay Platform is now PCI-DSS Level 1 compliant (upgraded from Level 2). "
        "SOC 2 Type II audit was completed in Q2 2024 (previously only Type I). "
        "GDPR data processing agreements are now available for EU customers. "
        "TLS minimum version updated to 1.3."
    ))

    _add_heading(doc, "Support SLA Updates — Effective Q1 2025")
    _add_para(doc, (
        "Support has been significantly expanded:\n\n"
        "• Starter: Email support, 48-hour response (unchanged).\n"
        "• Growth: Email + live chat support, 12-hour response (was 24-hour email only).\n"
        "• Business: Email + live chat + phone, 4-hour response, Monday–Friday 8 AM – 8 PM EST.\n"
        "• Enterprise: 24/7 phone and dedicated Slack channel, 1-hour response SLA.\n\n"
        "The self-service knowledge base URL has changed to support.novapay.io (was help.novapay.io)."
    ))

    _add_heading(doc, "Company Milestones (for About section)")
    _add_para(doc, (
        "Updated stats for use in collateral:\n"
        "• Now processing over $1.2 billion in payments annually (was $200M).\n"
        "• Customer base: 18,000+ businesses across North America and Europe (was 3,000).\n"
        "• Series B funding completed October 2024: $42M raised.\n"
        "• New EU headquarters opened in Dublin, Ireland (Q4 2024)."
    ))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


DEMO_OUTDATED_FILENAME = "NovaPay_Pro_Product_Brief_2022.docx"
DEMO_REFERENCE_FILENAME = "NovaPay_Internal_Q1_2025_Update.docx"
DEMO_CONTEXT_NOTE = (
    "The product has been rebranded from 'NovaPay Pro' to 'NovaPay Platform'. "
    "Pricing, feature set, integrations, SLA, and company stats have all been updated "
    "as of Q1 2025. All customer-facing copy must reflect the new name and current details."
)


def get_demo_files() -> tuple[bytes, bytes, str, str, str]:
    """
    Returns (outdated_bytes, reference_bytes, outdated_filename, reference_filename, context_note).
    """
    return (
        build_outdated_document(),
        build_reference_document(),
        DEMO_OUTDATED_FILENAME,
        DEMO_REFERENCE_FILENAME,
        DEMO_CONTEXT_NOTE,
    )
